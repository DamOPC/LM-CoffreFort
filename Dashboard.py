#Packages
import streamlit as st
import pandas as pd
import docx
import io
import openpyxl

st.set_page_config(layout='wide')

st.title('Génération automatique d\'AG')

df = pd.read_excel('AGAuto50.xlsx')
doc = docx.Document('edoc.docx')
bio = io.BytesIO()

#Infos lettre mission
choix_client = st.selectbox("Choisir un code client", options=df['CODE Client'])
date_cr = st.date_input("Selectionner la date d'edition du document")
mois = st.text_input("Entrer le mois de mise en place du ")
option = st.radio(["Option1", "Option2"])

#Fonction de generation des variables
def generate_var():
    df_client = df[df['CODE Client']==choix_client]
    _index = df_client.index.values[0]
    societe = df_client.loc[_index, 'Nom Client']
    adresse = df_client.loc[_index, 'Adresse']
    code_postal = str(df_client.loc[_index, 'Code Postal'])
    ville = df_client.loc[_index, 'Ville']
    siret = str(df_client.loc[_index, 'SIRET'])
    benefices = str(df_client.loc[_index, 'Resultat net'])
    nom_dir = df_client.loc[_index, 'Nom Dirigieant']
    parts_dir = str(df_client.loc[_index, 'Parts Dirigeant'])
    type = df_client.loc[_index, 'Titre']
    parts_totales = str(df_client.loc[_index, 'Relations.Nombre part entreprise'])
    date_pv = str(date_cr)
    
    replacements = {
    '%nom_societe%': societe,
    '%type%': type,
    '%capital%': '10000',
    '%rue%': adresse,
    '%ville%': ville,
    '%code_postal%': code_postal,
    '%siren%': siret,
    '%parts_totales%': parts_dir,
    '%benefices%': benefices,
    '%dirigeant%': nom_dir,
    '%part%': parts_dir,
    '%date%': date_pv,
    '%part_totales%': parts_totales,   
    }

    return replacements

#Fonction de modification du document
def generate_docx():
    replacements = generate_var()
    for paragraph in doc.paragraphs:
        for key in replacements:
            paragraph.text = paragraph.text.replace(key, replacements[key])

    doc.save(bio)
    return doc_replace
    
if st.button('Generez'): 
    #_index = df[df['CODE Client']=='10LYSCI'].index.values[0]
    #st.write('la ville est:', df[df['CODE Client']=='10LYSCI'])
    #st.write(_index)
    #st.write(df[df['CODE Client']=='10LYSCI'].loc[_index, 'Nom Client'])
    replacements = generate_var()
    for paragraph in doc.paragraphs:
        for key in replacements:
            paragraph.text = paragraph.text.replace(key, replacements[key])
    doc.save(bio)
    if doc:
        st.download_button(
            label="Cliquer ici pour télécharger la Lettre de Mission",
            data=bio.getvalue(),
            file_name="LM_CoffreFort_EDOC.docx",
            mime="docx"
        )


#df_client = df[df["CODE Client"]=='client_champ']
#_index = df_client.index.values[0]

#societe = df_client.loc[_index, 'Nom Client']
#st.write('la ville du client est:', ville)
